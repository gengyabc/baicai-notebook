import argparse
import csv
import json
import re
from pathlib import Path

from docx import Document

from .exceptions import TemplateGenError
from .task_paths import TaskPaths


SUPPORTED_TEMPLATE_TOKEN_PATTERN = re.compile(
    r"\{\{\s*[A-Za-z_][A-Za-z0-9_]*(?:\.[A-Za-z_][A-Za-z0-9_]*)?\s*\}\}|\{%\s*for\s+\w+\s+in\s+\w+\s*%\}|\{%\s*endfor\s*%\}"
)


def _extract_supported_tokens(text: str) -> list[str]:
    return SUPPORTED_TEMPLATE_TOKEN_PATTERN.findall(text)


def extract_template_placeholders(template_docx_path: str) -> list[dict[str, str]]:
    source = Path(template_docx_path)
    if not source.exists():
        raise TemplateGenError(f"Template file not found: {template_docx_path}")

    try:
        document = Document(str(source))
    except Exception as exc:
        raise TemplateGenError(f"Failed to read template DOCX: {template_docx_path}") from exc

    extracted: list[dict[str, str]] = []

    for para_idx, paragraph in enumerate(document.paragraphs):
        for token in _extract_supported_tokens(paragraph.text):
            extracted.append({"location": f"paragraphs[{para_idx}]", "placeholder": token})

    for table_idx, table in enumerate(document.tables):
        for row_idx, row in enumerate(table.rows):
            for col_idx, cell in enumerate(row.cells):
                for paragraph in cell.paragraphs:
                    for token in _extract_supported_tokens(paragraph.text):
                        extracted.append({
                            "location": f"tables[{table_idx}].rows[{row_idx}].cells[{col_idx}]",
                            "placeholder": token,
                        })

    if not extracted:
        raise TemplateGenError(
            "No supported placeholders found in template.docx. "
            "Run /generate-template or update template placeholders first."
        )

    return extracted


def rebuild_placeholders_from_template(template_docx_path: str, placeholders_output_path: str) -> str:
    placeholders = extract_template_placeholders(template_docx_path)
    
    # Deduplicate: keep only first occurrence of each unique placeholder
    seen: set[str] = set()
    deduped: list[dict[str, str]] = []
    for ph in placeholders:
        placeholder = ph["placeholder"]
        if placeholder not in seen:
            seen.add(placeholder)
            deduped.append(ph)
    
    output = Path(placeholders_output_path)
    output.parent.mkdir(parents=True, exist_ok=True)
    output.write_text(
        json.dumps({"placeholders": deduped}, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )
    return str(output)


def load_placeholder_description_source(json_path: str) -> list[dict[str, str]]:
    source = Path(json_path)
    if not source.exists():
        raise TemplateGenError(f"Input file not found: {json_path}")

    try:
        data = json.loads(source.read_text(encoding="utf-8"))
    except json.JSONDecodeError as exc:
        raise TemplateGenError(f"Invalid JSON: {json_path}") from exc

    placeholders = data.get("placeholders")
    if not isinstance(placeholders, list):
        raise TemplateGenError("Input JSON must contain a placeholders array")

    unique_rows: list[dict[str, str]] = []
    first_seen_descriptions: dict[str, str] = {}

    for item in placeholders:
        if not isinstance(item, dict):
            raise TemplateGenError("Each placeholders item must be an object")

        raw_placeholder = item.get("placeholder")
        if not isinstance(raw_placeholder, str) or raw_placeholder == "":
            raise TemplateGenError("Each placeholders item must include a non-empty placeholder")

        description = item.get("description") or ""
        if not isinstance(description, str):
            raise TemplateGenError("description must be a string when provided")

        if raw_placeholder in first_seen_descriptions:
            first_desc = first_seen_descriptions[raw_placeholder]
            if first_desc and description and first_desc != description:
                raise TemplateGenError(
                    f"Conflicting non-empty descriptions for placeholder: {raw_placeholder}"
                )
            continue

        first_seen_descriptions[raw_placeholder] = description
        unique_rows.append({"placeholder": raw_placeholder, "description": description})

    return unique_rows


def export_placeholder_csv(
    input_path: str,
    output_path: str,
    *,
    edit: bool = False,
    template_docx_path: str | None = None,
    placeholders_output_path: str | None = None,
) -> str:
    source_path = input_path

    if edit:
        if not template_docx_path:
            raise TemplateGenError("edit mode requires template_docx_path")
        if not placeholders_output_path:
            raise TemplateGenError("edit mode requires placeholders_output_path")
        source_path = rebuild_placeholders_from_template(
            template_docx_path,
            placeholders_output_path,
        )

    rows = load_placeholder_description_source(source_path)
    output = Path(output_path)
    output.parent.mkdir(parents=True, exist_ok=True)

    with output.open("w", encoding="utf-8", newline="") as f:
        writer = csv.DictWriter(f, fieldnames=["placeholder", "description"])
        writer.writeheader()
        writer.writerows(rows)

    return str(output)


def main() -> None:
    parser = argparse.ArgumentParser(
        description="Export reviewed placeholders JSON into a minimal placeholder/description CSV"
    )
    parser.add_argument(
        "--edit",
        action="store_true",
        help="Rebuild placeholders from current template.docx before CSV export",
    )
    parser.add_argument("--input", help="Reviewed placeholders JSON file (optional)")
    parser.add_argument("--output", help="Output CSV file (optional)")
    parser.add_argument("--template", help="Current template DOCX file (optional)")
    args = parser.parse_args()

    if args.input and args.output:
        output = export_placeholder_csv(
            args.input,
            args.output,
            edit=args.edit,
            template_docx_path=args.template if args.edit else None,
            placeholders_output_path=args.input if args.edit else None,
        )
    else:
        task_paths = TaskPaths.get_current()
        output = export_placeholder_csv(
            str(task_paths.placeholders_json),
            str(task_paths.descriptions_csv),
            edit=args.edit,
            template_docx_path=str(task_paths.template_docx) if args.edit else None,
            placeholders_output_path=str(task_paths.placeholders_json) if args.edit else None,
        )

    print(f"Placeholder CSV exported to: {output}")


if __name__ == "__main__":
    main()
