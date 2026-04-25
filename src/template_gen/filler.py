from copy import deepcopy
import os
import re
import tempfile

from docx import Document
from docxtpl import DocxTemplate
from jinja2 import Environment, StrictUndefined

from .exceptions import FillError
from .schemas import CoordinateMapping, DocumentStructure, FieldInfo, PlaceholderMapping


LOOP_ROW_PATTERN = re.compile(r"\{\%\s*for\s+(\w+)\s+in\s+(\w+)\s*\%\}")
LOOP_FIELD_PATTERN = re.compile(r"\{\{\s*(\w+)\.(\w+)\s*\}\}")


def detect_merged_cells(table) -> dict:
    """Detect merged cells in a table. Returns dict mapping (row, col) -> primary cell coords."""
    merged_map = {}
    seen_cells = {}

    for row_idx, row in enumerate(table.rows):
        for col_idx, cell in enumerate(row.cells):
            tc = cell._tc
            if tc in seen_cells:
                merged_map[(row_idx, col_idx)] = seen_cells[tc]
            else:
                seen_cells[tc] = (row_idx, col_idx)

    return merged_map


def _copy_run_style(source_run, target_run) -> None:
    if source_run is None or source_run._element.rPr is None:
        return

    target_rpr = target_run._element.get_or_add_rPr()
    for child in list(target_rpr):
        target_rpr.remove(child)
    for child in source_run._element.rPr:
        target_rpr.append(deepcopy(child))


def _find_style_source_run(table, row_idx: int, col_idx: int):
    candidates = []

    row = table.rows[row_idx]
    for other_col, cell in enumerate(row.cells):
        if other_col != col_idx:
            candidates.extend(cell.paragraphs[0].runs)

    for other_row, table_row in enumerate(table.rows):
        if other_row != row_idx and col_idx < len(table_row.cells):
            candidates.extend(table_row.cells[col_idx].paragraphs[0].runs)

    for run in candidates:
        if run.text or run._element.rPr is not None:
            return run

    return None


def _apply_rpr_to_run(run, rpr_element) -> None:
    if rpr_element is None:
        return
    target_rpr = run._element.get_or_add_rPr()
    for child in list(target_rpr):
        target_rpr.remove(child)
    for child in rpr_element:
        target_rpr.append(deepcopy(child))


def set_cell_text_keep_style(cell, text: str, style_source_run=None, style_rpr=None) -> None:
    """Set cell text while preserving or borrowing run formatting."""
    paragraph = cell.paragraphs[0]

    if paragraph.runs:
        run = paragraph.runs[0]
        run.text = text
        for extra_run in paragraph.runs[1:]:
            extra_run.text = ""
        if style_rpr is not None:
            _apply_rpr_to_run(run, style_rpr)
    else:
        run = paragraph.add_run(text)
        if style_rpr is not None:
            _apply_rpr_to_run(run, style_rpr)
        elif style_source_run is not None:
            _copy_run_style(style_source_run, run)


def set_paragraph_text_keep_style(paragraph, text: str) -> None:
    """Set paragraph text while preserving existing run formatting."""
    if paragraph.runs:
        run = paragraph.runs[0]
        run.text = text
        for extra_run in paragraph.runs[1:]:
            extra_run.text = ""
    else:
        paragraph.add_run(text)


def set_cell_text_keep_basic_style(cell, text: str, style_source_run=None) -> None:
    """Backward-compatible alias for style-preserving cell replacement."""
    set_cell_text_keep_style(cell, text, style_source_run=style_source_run)


def _collect_loop_table_specs(doc: Document) -> list[dict]:
    specs = []

    for table_index, table in enumerate(doc.tables):
        for row_index, row in enumerate(table.rows):
            row_text = " ".join(cell.text for cell in row.cells)
            loop_match = LOOP_ROW_PATTERN.search(row_text)
            if not loop_match:
                continue

            loop_var = loop_match.group(1)
            list_name = loop_match.group(2)
            field_map: dict[int, str] = {}
            style_rprs: dict[int, object] = {}

            for col_index, cell in enumerate(row.cells):
                cell_match = LOOP_FIELD_PATTERN.search(cell.text)
                if not cell_match or cell_match.group(1) != loop_var:
                    continue

                field_map[col_index] = cell_match.group(2)
                run = cell.paragraphs[0].runs[0] if cell.paragraphs[0].runs else None
                if run is not None and run._element.rPr is not None:
                    style_rprs[col_index] = deepcopy(run._element.rPr)
                else:
                    style_rprs[col_index] = None

            if field_map:
                specs.append(
                    {
                        "table_index": table_index,
                        "row_index": row_index,
                        "list_name": list_name,
                        "field_map": field_map,
                        "style_rprs": style_rprs,
                    }
                )

    return specs


def _blank_table_rows(doc: Document, specs: list[dict]) -> None:
    for spec in specs:
        table = doc.tables[spec["table_index"]]
        if spec["row_index"] >= len(table.rows):
            continue

        row = table.rows[spec["row_index"]]
        for cell in row.cells:
            cell.text = ""


def _fill_loop_tables(doc: Document, data: dict, specs: list[dict]) -> None:
    for spec in specs:
        table = doc.tables[spec["table_index"]]
        items = data.get(spec["list_name"], [])
        if not isinstance(items, list):
            continue

        start_row = spec["row_index"]
        field_map = spec["field_map"]
        style_rprs = spec["style_rprs"]

        for offset, item in enumerate(items):
            row_index = start_row + offset
            if row_index >= len(table.rows):
                break

            row = table.rows[row_index]
            item_data = item if isinstance(item, dict) else {}

            for col_index, field_name in field_map.items():
                value = item_data.get(field_name, "")
                if value is None:
                    value = ""
                set_cell_text_keep_style(
                    row.cells[col_index],
                    str(value),
                    style_rpr=style_rprs.get(col_index),
                )

        for row_index in range(start_row + len(items), len(table.rows)):
            row = table.rows[row_index]
            for cell in row.cells:
                if cell.text.strip():
                    cell.text = ""


def _snake_case(text: str) -> str:
    """Convert Chinese/English text to snake_case placeholder name."""
    import re
    text = text.strip()
    # Remove common filler words
    text = re.sub(r'[（(].*?[)）]', '', text)
    text = text.replace(' ', '').replace('　', '')
    # Simple pinyin-like transliteration for common Chinese chars isn't available,
    # so we'll use a direct mapping for common terms and fallback to generic names.
    return text


def _extract_header_text(table, row_idx: int, col_idx: int, merged_map: dict, header_row_set: set = None) -> str:
    """Try to extract column header text from header rows above the given cell."""
    for hr in range(row_idx - 1, -1, -1):
        if hr >= len(table.rows):
            continue
        # Skip non-header rows if header_row_set is provided
        if header_row_set is not None and hr not in header_row_set:
            continue
        row = table.rows[hr]
        if col_idx >= len(row.cells):
            continue
        cell = row.cells[col_idx]
        primary = merged_map.get((hr, col_idx), (hr, col_idx))
        if primary != (hr, col_idx):
            continue
        txt = cell.text.strip()
        if txt:
            return txt
    return ""


def _extract_row_label(table, row_idx: int, merged_map: dict) -> str:
    """Try to extract row label from the first column of the given row."""
    if len(table.rows[row_idx].cells) == 0:
        return ""
    cell = table.rows[row_idx].cells[0]
    txt = cell.text.strip()
    return txt


def _build_auto_placeholder_internal(table, row_idx: int, col_idx: int, merged_map: dict, existing: set, header_row_set: set) -> str:
    """Build an auto-generated placeholder name for an empty cell."""
    import re
    header = _extract_header_text(table, row_idx, col_idx, merged_map, header_row_set)
    row_label = _extract_row_label(table, row_idx, merged_map)

    # Common Chinese term mappings
    term_map = {
        '获奖时间': 'award_time',
        '获奖种类': 'award_category',
        '获奖等级': 'award_level',
        '授奖部门': 'award_issuer',
        '教材名称': 'textbook_name',
        '编写': 'writing_staff',
        '第一主编': 'first_editor',
        '第一主编所属单位': 'first_editor_unit',
        '申报单位联系方式': 'applicant_contact',
        '适用专业代码及名称': 'applicable_major',
        '编写团队成员': 'writing_team',
        '对应课程性质': 'course_nature',
        '对应领域': 'corresponding_field',
        '书号': 'isbn',
        '版次': 'version',
        '出版时间': 'publish_date',
        '初版时间': 'initial_publish_date',
        '印数': 'print_count',
        '累计发行量': 'cumulative_circulation',
        '单位名称': 'unit_name',
        '主管部门': 'supervising_department',
        '联系人': 'contact_person',
        '联系电话': 'contact_phone',
        '电子邮箱': 'email',
        '通讯地址': 'postal_address',
        '邮政编码': 'postal_code',
    }

    # Normalize header for matching (remove spaces and newlines)
    header_normalized = header.replace(' ', '').replace('　', '').replace('\n', '').replace('\r', '')

    # Try to match header
    name = None
    for key, val in term_map.items():
        if key in header_normalized:
            name = val
            break

    if not name:
        # Clean header for a basic name
        name = re.sub(r'[^\w\u4e00-\u9fff]+', '_', header_normalized).strip('_')
        if not name:
            name = f"cell_{col_idx}"

    # Detect award rows and use indexed names like award_1_category
    if '教材获奖情况' in row_label or '获奖' in row_label:
        # Find the header row for this section using pre-scanned header_row_set
        header_row_idx = None
        for hr in range(row_idx, -1, -1):
            hr_label = _extract_row_label(table, hr, merged_map)
            if '教材获奖情况' in hr_label or '获奖' in hr_label:
                if hr in header_row_set:
                    header_row_idx = hr
                    break
        
        # Count data rows after header
        if header_row_idx is not None:
            award_index = row_idx - header_row_idx
            # Use simpler naming: award_1_category instead of award_1_award_category
            # Remove 'award_' prefix from name if already present
            simple_name = name.replace('award_', '')
            base = f"award_{award_index}_{simple_name}"
        else:
            base = name
    else:
        base = name

    candidate = f"{{{{ {base} }}}}"
    counter = 1
    while candidate in existing:
        candidate = f"{{{{ {base}_{counter} }}}}"
        counter += 1
    return candidate


def _build_auto_placeholder(table, row_idx: int, col_idx: int, merged_map: dict, existing: set) -> str:
    """Legacy wrapper for backward compatibility."""
    return _build_auto_placeholder_internal(table, row_idx, col_idx, merged_map, existing, set())


def generate_template(
    input_path: str,
    output_path: str,
    mappings: list[CoordinateMapping],
) -> DocumentStructure:
    doc = Document(input_path)
    
    merged_maps = {}
    for t_idx, table in enumerate(doc.tables):
        merged_maps[t_idx] = detect_merged_cells(table)
    
    # Pre-scan to identify header rows based on original content (BEFORE any modifications)
    header_rows = {}  # table_idx -> set of header row indices
    for t_idx, table in enumerate(doc.tables):
        header_rows[t_idx] = set()
        for r_idx, row in enumerate(table.rows):
            # A row is considered a header if most cells have non-empty text
            non_empty = sum(1 for cc in row.cells if cc.text.strip())
            if non_empty > len(row.cells) // 2:
                header_rows[t_idx].add(r_idx)
    
    processed_cells = set()
    existing_placeholders = set()
    
    for mapping in mappings:
        if not mapping.is_empty:
            continue
        existing_placeholders.add(mapping.placeholder)

        if mapping.target_type == "paragraph":
            if mapping.paragraph_index is None or mapping.paragraph_index >= len(doc.paragraphs):
                continue
            paragraph = doc.paragraphs[mapping.paragraph_index]
            if paragraph.text.strip():
                if paragraph.runs:
                    paragraph.runs[-1].text += " " + mapping.placeholder
                else:
                    paragraph.add_run(" " + mapping.placeholder)
                continue
            set_paragraph_text_keep_style(paragraph, mapping.placeholder)
            continue
        
        table_idx = mapping.table_index
        row_idx = mapping.row_index
        col_idx = mapping.col_index
        if table_idx is None or row_idx is None or col_idx is None:
            continue
        
        cell_key = (table_idx, row_idx, col_idx)
        if cell_key in processed_cells:
            continue
        
        table = doc.tables[table_idx]
        row = table.rows[row_idx]
        cell = row.cells[col_idx]

        if cell.text.strip():
            continue

        merged_map = merged_maps.get(table_idx, {})
        primary_key = merged_map.get((row_idx, col_idx), (row_idx, col_idx))
        
        if primary_key != (row_idx, col_idx):
            continue
        
        style_source_run = None
        if not cell.paragraphs[0].runs:
            style_source_run = _find_style_source_run(table, row_idx, col_idx)

        set_cell_text_keep_style(cell, mapping.placeholder, style_source_run=style_source_run)
        processed_cells.add(cell_key)
    
    # Pass header row info to placeholder builder
    def _build_auto_placeholder_with_headers(table, row_idx, col_idx, merged_map, existing, header_row_set):
        return _build_auto_placeholder_internal(table, row_idx, col_idx, merged_map, existing, header_row_set)
    
    # Auto-fill remaining empty primary cells
    for t_idx, table in enumerate(doc.tables):
        merged_map = merged_maps.get(t_idx, {})
        header_row_set = header_rows.get(t_idx, set())
        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                cell_key = (t_idx, r_idx, c_idx)
                if cell_key in processed_cells:
                    continue
                
                primary_key = merged_map.get((r_idx, c_idx), (r_idx, c_idx))
                if primary_key != (r_idx, c_idx):
                    continue
                
                if cell.text.strip():
                    continue
                
                # Skip the very first column if it's acting as a row label
                if c_idx == 0:
                    continue
                
                auto_placeholder = _build_auto_placeholder_with_headers(
                    table, r_idx, c_idx, merged_map, existing_placeholders, header_row_set
                )
                existing_placeholders.add(auto_placeholder)
                
                style_source_run = None
                if not cell.paragraphs[0].runs:
                    style_source_run = _find_style_source_run(table, r_idx, c_idx)
                
                set_cell_text_keep_style(cell, auto_placeholder, style_source_run=style_source_run)
                processed_cells.add(cell_key)
    
    doc.save(output_path)
    
    from .parser import parse_document
    return parse_document(output_path, save_json=False)


def fill_template(
    template_path: str,
    data: dict,
    output_path: str,
) -> str:
    if not os.path.exists(template_path):
        raise FillError(f"Template file not found: {template_path}")
    
    try:
        template_doc = Document(template_path)
        loop_specs = _collect_loop_table_specs(template_doc)

        working_template_path = template_path
        temp_path = None
        if loop_specs:
            _blank_table_rows(template_doc, loop_specs)
            tmp = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
            temp_path = tmp.name
            tmp.close()
            template_doc.save(temp_path)
            working_template_path = temp_path

        tpl = DocxTemplate(working_template_path)
        tpl.render(data, jinja_env=Environment(undefined=StrictUndefined))
        tpl.save(output_path)

        if loop_specs:
            rendered_doc = Document(output_path)
            _fill_loop_tables(rendered_doc, data, loop_specs)
            rendered_doc.save(output_path)

        if temp_path and os.path.exists(temp_path):
            os.unlink(temp_path)

        return output_path
    except Exception as e:
        if "temp_path" in locals() and temp_path and os.path.exists(temp_path):
            os.unlink(temp_path)
        raise FillError(f"Fill failed: {str(e)}")


def _sanitize_filename_component(value: str) -> str:
    sanitized = value.strip().replace("..", "_")
    for separator in (os.sep, os.altsep):
        if separator:
            sanitized = sanitized.replace(separator, "_")
    return sanitized or "output"


def _build_output_path(output_dir: str, output_filename: str) -> str:
    output_dir_abs = os.path.abspath(output_dir)
    output_path = os.path.abspath(os.path.join(output_dir_abs, output_filename))

    if os.path.commonpath([output_dir_abs, output_path]) != output_dir_abs:
        raise FillError(f"Unsafe output filename: {output_filename}")

    return output_path


def batch_fill(
    template_path: str,
    data_records: list[dict],
    output_dir: str,
    filename_pattern: str = "output_{index}.docx",
) -> list[str]:
    if not os.path.exists(template_path):
        raise FillError(f"Template file not found: {template_path}")
    
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)
    
    outputs = []
    failed = []
    
    for i, data in enumerate(data_records):
        output_filename = filename_pattern.replace("{index}", str(i + 1))
        if "{name}" in filename_pattern and "name" in data:
            output_filename = output_filename.replace("{name}", _sanitize_filename_component(str(data["name"])))
        output_path = _build_output_path(output_dir, output_filename)

        try:
            fill_template(template_path, data, output_path)
            outputs.append(output_path)
        except FillError as e:
            failed.append({"index": i, "error": str(e)})
    
    if failed:
        print(f"Warning: {len(failed)} records failed to fill")
    
    return outputs
