import json
from pathlib import Path

from .exceptions import TemplateGenError
from .schemas import CoordinateMapping
from .task_paths import TaskPaths
from .filler import generate_template as _generate_template


def parse_location(location: str) -> dict:
    if location.startswith("paragraphs[") and location.endswith("]"):
        try:
            return {"paragraphs": int(location[len("paragraphs[") : -1])}
        except ValueError as exc:
            raise TemplateGenError(f"Invalid location index: {location}") from exc

    parts = location.split('.')
    expected_keys = ("tables", "rows", "cells")
    result = {}

    if len(parts) != len(expected_keys):
        raise TemplateGenError(f"Unsupported location: {location}")

    for part, expected_key in zip(parts, expected_keys):
        if "[" not in part or not part.endswith("]"):
            raise TemplateGenError(f"Unsupported location: {location}")

        key, raw_index = part.split("[", 1)
        if key != expected_key:
            raise TemplateGenError(f"Unsupported location: {location}")

        try:
            result[key] = int(raw_index[:-1])
        except ValueError as exc:
            raise TemplateGenError(f"Invalid location index: {location}") from exc

    return result


def load_placeholders_json(json_path: str) -> list[CoordinateMapping]:
    with open(json_path, 'r', encoding='utf-8') as f:
        data = json.load(f)
    
    mappings = []
    for placeholder in data['placeholders']:
        loc = parse_location(placeholder['location'])
        if 'paragraphs' in loc:
            mapping = CoordinateMapping(
                placeholder=placeholder['placeholder'],
                is_empty=True,
                target_type='paragraph',
                paragraph_index=loc['paragraphs'],
            )
        else:
            mapping = CoordinateMapping(
                placeholder=placeholder['placeholder'],
                is_empty=True,
                target_type='cell',
                table_index=loc.get('tables', 0),
                row_index=loc.get('rows', 0),
                col_index=loc.get('cells', 0),
            )
        mappings.append(mapping)
    
    return mappings


def generate_template_from_json(
    source_docx: str,
    placeholders_json: str,
    output_path: str,
) -> None:
    mappings = load_placeholders_json(placeholders_json)
    document_structure = _generate_template(source_docx, output_path, mappings)
    
    # After template generation, sync placeholders.json with actual template content
    # This captures both explicit and auto-generated placeholders
    _sync_placeholders_with_template(output_path, placeholders_json)
    
    print(f"Template saved to: {output_path}")


def _sync_placeholders_with_template(template_path: str, placeholders_json_path: str) -> None:
    """Update placeholders.json to reflect all placeholders in the generated template."""
    import re
    from docx import Document
    
    doc = Document(template_path)
    seen: set[str] = set()
    synced: list[dict] = []
    
    # Load existing placeholders.json to preserve context info
    with open(placeholders_json_path, 'r', encoding='utf-8') as f:
        existing_data = json.load(f)
    
    existing_by_placeholder: dict[str, dict] = {}
    for item in existing_data.get('placeholders', []):
        ph = item.get('placeholder', '')
        if ph and ph not in existing_by_placeholder:
            existing_by_placeholder[ph] = item
    
    placeholder_pattern = re.compile(r'\{\{\s*([A-Za-z_][A-Za-z0-9_]*)\s*\}\}')
    
    # Scan paragraphs
    for para_idx, paragraph in enumerate(doc.paragraphs):
        for match in placeholder_pattern.finditer(paragraph.text):
            ph = f"{{{{ {match.group(1)} }}}}"
            if ph not in seen:
                seen.add(ph)
                existing = existing_by_placeholder.get(ph, {})
                synced.append({
                    "location": f"paragraphs[{para_idx}]",
                    "placeholder": ph,
                    "context": existing.get("context", ""),
                    "field_path": existing.get("field_path", match.group(1)),
                })
    
    # Scan tables
    for table_idx, table in enumerate(doc.tables):
        for row_idx, row in enumerate(table.rows):
            for col_idx, cell in enumerate(row.cells):
                for match in placeholder_pattern.finditer(cell.text):
                    ph = f"{{{{ {match.group(1)} }}}}"
                    if ph not in seen:
                        seen.add(ph)
                        existing = existing_by_placeholder.get(ph, {})
                        synced.append({
                            "location": f"tables[{table_idx}].rows[{row_idx}].cells[{col_idx}]",
                            "placeholder": ph,
                            "context": existing.get("context", ""),
                            "field_path": existing.get("field_path", match.group(1)),
                        })
    
    # Write updated placeholders.json
    with open(placeholders_json_path, 'w', encoding='utf-8') as f:
        json.dump({"placeholders": synced}, f, ensure_ascii=False, indent=2)


if __name__ == '__main__':
    import argparse
    
    parser = argparse.ArgumentParser(description='Generate Jinja template from DOCX')
    parser.add_argument('--source', help='Source DOCX file (optional, uses current task if not provided)')
    parser.add_argument('--placeholders', help='Placeholders JSON file (optional)')
    parser.add_argument('--output', help='Output template DOCX file (optional)')
    
    args = parser.parse_args()
    
    if args.source and args.placeholders and args.output:
        Path(args.output).parent.mkdir(parents=True, exist_ok=True)
        generate_template_from_json(args.source, args.placeholders, args.output)
    else:
        task_paths = TaskPaths.get_current()
        generate_template_from_json(
            str(task_paths.input_docx),
            str(task_paths.placeholders_json),
            str(task_paths.template_docx),
        )
