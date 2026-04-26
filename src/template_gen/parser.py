import json
from dataclasses import asdict
from datetime import datetime
from pathlib import Path
from .schemas import DocumentStructure, ParagraphInfo, TableInfo, RowInfo, CellInfo, StyleInfo
from .exceptions import ParseError, TemplateGenError
from .task_paths import TaskPaths, TaskState
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from typing import Optional


def _extract_style(paragraph) -> Optional[StyleInfo]:
    style_info = StyleInfo()
    has_style = False
    
    if paragraph.alignment:
        alignment_map = {
            WD_ALIGN_PARAGRAPH.LEFT: "left",
            WD_ALIGN_PARAGRAPH.CENTER: "center",
            WD_ALIGN_PARAGRAPH.RIGHT: "right",
            WD_ALIGN_PARAGRAPH.JUSTIFY: "justify",
        }
        style_info.alignment = alignment_map.get(paragraph.alignment)
        has_style = True
    
    for run in paragraph.runs:
        if run.font.name:
            style_info.font_name = run.font.name
            has_style = True
        if run.font.size:
            style_info.font_size = run.font.size.pt
            has_style = True
        if run.bold:
            style_info.bold = True
            has_style = True
        if run.italic:
            style_info.italic = True
            has_style = True
        break
    
    return style_info if has_style else None


def parse_document(file_path: str, output_path: Optional[str] = None, save_json: bool = True) -> DocumentStructure:
    try:
        doc = Document(file_path)
    except Exception as e:
        raise ParseError(file_path, str(e))
    
    paragraphs = []
    for i, para in enumerate(doc.paragraphs):
        style = _extract_style(para)
        paragraphs.append(ParagraphInfo(
            index=i,
            text=para.text,
            style=style
        ))
    
    tables = []
    for t_idx, table in enumerate(doc.tables):
        rows = []
        for r_idx, row in enumerate(table.rows):
            cells = []
            for c_idx, cell in enumerate(row.cells):
                para = cell.paragraphs[0] if cell.paragraphs else None
                text = para.text if para else ""
                is_empty = text.strip() == ""
                style = _extract_style(para) if para else None
                cells.append(CellInfo(
                    row_index=r_idx,
                    col_index=c_idx,
                    text=text,
                    is_empty=is_empty,
                    style=style
                ))
            rows.append(RowInfo(index=r_idx, cells=cells))
        tables.append(TableInfo(index=t_idx, rows=rows))
    
    styles = {}
    for style in doc.styles:
        if style.type == 1:
            styles[style.name] = StyleInfo(
                font_name=style.font.name if style.font else None,
                font_size=style.font.size.pt if style.font and style.font.size else None,
            )
    
    structure = DocumentStructure(paragraphs=paragraphs, tables=tables, styles=styles)
    
    if save_json and output_path:
        Path(output_path).parent.mkdir(parents=True, exist_ok=True)
        with open(output_path, "w", encoding="utf-8") as f:
            json.dump(asdict(structure), f, ensure_ascii=False, indent=2)
    
    return structure


def print_table_coordinates(file_path: str) -> str:
    structure = parse_document(file_path)
    output_lines = []
    
    for table in structure.tables:
        output_lines.append(f"Table {table.index}:")
        for row in table.rows:
            for cell in row.cells:
                status = "empty" if cell.is_empty else f"'{cell.text[:20]}'"
                output_lines.append(f"  [{row.index},{cell.col_index}] {status}")
    
    return "\n".join(output_lines)


if __name__ == '__main__':
    import argparse
    
    parser = argparse.ArgumentParser(description='Parse Word document structure')
    parser.add_argument('docx_file', nargs='?', help='Path to .docx file (optional, finds latest if not provided)')
    
    args = parser.parse_args()
    
    if args.docx_file:
        task_paths = TaskPaths.create_or_next_version(args.docx_file)
        structure = parse_document(str(task_paths.input_docx), str(task_paths.parsed_json))
        print(f"Parsed: {args.docx_file}")
    else:
        docx_path = TaskPaths.find_latest_input_docx()
        if docx_path:
            task = docx_path.stem
            version = TaskPaths._get_latest_version(task) or 1
            task_paths = TaskPaths(task=task, version=version)
            TaskPaths._write_state(TaskState(
                task=task,
                version=version,
                updated_at=datetime.now().isoformat(timespec="seconds"),
            ))
            structure = parse_document(str(docx_path), str(task_paths.parsed_json), save_json=True)
            print(f"Found latest: {docx_path}")
        else:
            raise TemplateGenError("No docx files found in .temp/*/input/")
    
    print(f"Task: {task_paths.task}, Version: {task_paths.version}")
    print(f"Output: {task_paths.parsed_json}")
    print(f"Tables: {len(structure.tables)}, Paragraphs: {len(structure.paragraphs)}")