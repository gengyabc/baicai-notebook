import csv
import json
from pathlib import Path

import pytest
from docx import Document

from template_gen.exceptions import TemplateGenError
from template_gen.export_placeholder_csv import export_placeholder_csv, main as export_main
from template_gen.import_placeholder_csv import import_placeholder_csv, main as import_main


def _write_json(path: Path, payload: dict) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(json.dumps(payload, ensure_ascii=False), encoding="utf-8")


def _read_csv_rows(path: Path) -> list[dict[str, str]]:
    with path.open("r", encoding="utf-8", newline="") as f:
        return list(csv.DictReader(f))


class TestExportPlaceholderCsv:
    def test_export_writes_minimal_columns_with_blank_description(self, tmp_path: Path):
        source = tmp_path / "reviewed.json"
        out_csv = tmp_path / "out" / "placeholder_descriptions.csv"
        _write_json(
            source,
            {
                "placeholders": [
                    {"location": "tables[0].rows[0].cells[1]", "placeholder": "{{ project_name }}"},
                    {"location": "tables[0].rows[1].cells[1]", "placeholder": "{{ contact_name }}"},
                ]
            },
        )

        export_placeholder_csv(str(source), str(out_csv))

        assert out_csv.exists()
        assert out_csv.read_text(encoding="utf-8").splitlines()[0] == "placeholder,description"
        rows = _read_csv_rows(out_csv)
        assert rows == [
            {"placeholder": "{{ project_name }}", "description": ""},
            {"placeholder": "{{ contact_name }}", "description": ""},
        ]

    def test_export_deduplicates_by_first_occurrence_order(self, tmp_path: Path):
        source = tmp_path / "reviewed.json"
        out_csv = tmp_path / "placeholder_descriptions.csv"
        _write_json(
            source,
            {
                "placeholders": [
                    {"placeholder": "{{ a }}", "description": "字段A"},
                    {"placeholder": "{{ b }}", "description": "字段B"},
                    {"placeholder": "{{ a }}", "description": "字段A"},
                ]
            },
        )

        export_placeholder_csv(str(source), str(out_csv))

        assert _read_csv_rows(out_csv) == [
            {"placeholder": "{{ a }}", "description": "字段A"},
            {"placeholder": "{{ b }}", "description": "字段B"},
        ]

    def test_export_rejects_conflicting_non_empty_descriptions(self, tmp_path: Path):
        source = tmp_path / "reviewed.json"
        out_csv = tmp_path / "placeholder_descriptions.csv"
        _write_json(
            source,
            {
                "placeholders": [
                    {"placeholder": "{{ project_name }}", "description": "项目名称"},
                    {"placeholder": "{{ project_name }}", "description": "项目名"},
                ]
            },
        )

        with pytest.raises(TemplateGenError):
            export_placeholder_csv(str(source), str(out_csv))

    @pytest.mark.parametrize(
        "payload",
        [
            {},
            {"placeholders": [{"location": "tables[0].rows[0].cells[0]"}]},
            {"placeholders": [{"placeholder": ""}]},
            {"placeholders": "not-a-list"},
        ],
    )
    def test_export_rejects_invalid_input_shape(self, tmp_path: Path, payload: dict):
        source = tmp_path / "reviewed.json"
        out_csv = tmp_path / "placeholder_descriptions.csv"
        _write_json(source, payload)

        with pytest.raises(TemplateGenError):
            export_placeholder_csv(str(source), str(out_csv))

    def test_export_cli_creates_output_parent_directory(self, tmp_path: Path, monkeypatch):
        source = tmp_path / "reviewed.json"
        out_csv = tmp_path / "nested" / "dir" / "placeholder_descriptions.csv"
        _write_json(source, {"placeholders": [{"placeholder": "{{ project_name }}"}]})

        monkeypatch.setattr(
            "sys.argv",
            [
                "template_gen.export_placeholder_csv",
                "--input",
                str(source),
                "--output",
                str(out_csv),
            ],
        )
        export_main()

        assert out_csv.exists()

    def test_export_edit_rebuilds_placeholders_and_overwrites_csv(self, tmp_path: Path):
        placeholders_json = tmp_path / "temp" / "placeholders.json"
        out_csv = tmp_path / "output" / "descriptions.csv"
        template_docx = tmp_path / "output" / "template.docx"

        _write_json(
            placeholders_json,
            {
                "placeholders": [
                    {"location": "paragraphs[0]", "placeholder": "{{ stale_field }}"},
                ]
            },
        )

        doc = Document()
        doc.add_paragraph("标题：{{ project_name }}")
        table = doc.add_table(rows=3, cols=1)
        table.cell(0, 0).text = "{{ contact_name }}"
        table.cell(1, 0).text = "{{ contact_name }}"
        table.cell(2, 0).text = "{{ project_name }}"
        template_docx.parent.mkdir(parents=True, exist_ok=True)
        doc.save(template_docx)

        export_placeholder_csv(
            str(placeholders_json),
            str(out_csv),
            edit=True,
            template_docx_path=str(template_docx),
            placeholders_output_path=str(placeholders_json),
        )

        refreshed = json.loads(placeholders_json.read_text(encoding="utf-8"))
        assert refreshed == {
            "placeholders": [
                {"location": "paragraphs[0]", "placeholder": "{{ project_name }}"},
                {"location": "tables[0].rows[0].cells[0]", "placeholder": "{{ contact_name }}"},
                {"location": "tables[0].rows[1].cells[0]", "placeholder": "{{ contact_name }}"},
                {"location": "tables[0].rows[2].cells[0]", "placeholder": "{{ project_name }}"},
            ]
        }
        assert _read_csv_rows(out_csv) == [
            {"placeholder": "{{ project_name }}", "description": ""},
            {"placeholder": "{{ contact_name }}", "description": ""},
        ]

    def test_export_edit_fails_when_template_missing(self, tmp_path: Path):
        placeholders_json = tmp_path / "temp" / "placeholders.json"
        out_csv = tmp_path / "output" / "descriptions.csv"
        _write_json(placeholders_json, {"placeholders": [{"placeholder": "{{ a }}"}]})

        with pytest.raises(TemplateGenError):
            export_placeholder_csv(
                str(placeholders_json),
                str(out_csv),
                edit=True,
                template_docx_path=str(tmp_path / "output" / "template.docx"),
                placeholders_output_path=str(placeholders_json),
            )

    def test_export_edit_fails_when_template_has_no_supported_placeholders(self, tmp_path: Path):
        placeholders_json = tmp_path / "temp" / "placeholders.json"
        out_csv = tmp_path / "output" / "descriptions.csv"
        template_docx = tmp_path / "output" / "template.docx"
        _write_json(placeholders_json, {"placeholders": [{"placeholder": "{{ a }}"}]})

        doc = Document()
        doc.add_paragraph("没有占位符")
        template_docx.parent.mkdir(parents=True, exist_ok=True)
        doc.save(template_docx)

        with pytest.raises(TemplateGenError):
            export_placeholder_csv(
                str(placeholders_json),
                str(out_csv),
                edit=True,
                template_docx_path=str(template_docx),
                placeholders_output_path=str(placeholders_json),
            )


class TestImportPlaceholderCsv:
    def test_import_writes_minimal_json_in_csv_row_order(self, tmp_path: Path):
        source_csv = tmp_path / "edited.csv"
        out_json = tmp_path / "out" / "placeholder_descriptions.json"
        source_csv.write_text(
            "placeholder,description\n{{ b }},字段B\n{{ a }},字段A\n",
            encoding="utf-8",
        )

        import_placeholder_csv(str(source_csv), str(out_json))

        payload = json.loads(out_json.read_text(encoding="utf-8"))
        assert payload == {
            "placeholders": [
                {"placeholder": "{{ b }}", "description": "字段B"},
                {"placeholder": "{{ a }}", "description": "字段A"},
            ]
        }

    @pytest.mark.parametrize(
        "csv_text",
        [
            "description,placeholder\n字段A,{{ a }}\n",
            "placeholder\n{{ a }}\n",
            "placeholder,description,extra\n{{ a }},字段A,x\n",
        ],
    )
    def test_import_rejects_invalid_headers(self, tmp_path: Path, csv_text: str):
        source_csv = tmp_path / "edited.csv"
        out_json = tmp_path / "placeholder_descriptions.json"
        source_csv.write_text(csv_text, encoding="utf-8")

        with pytest.raises(TemplateGenError):
            import_placeholder_csv(str(source_csv), str(out_json))

    def test_import_rejects_empty_or_duplicate_placeholder(self, tmp_path: Path):
        source_csv = tmp_path / "edited.csv"
        out_json = tmp_path / "placeholder_descriptions.json"
        source_csv.write_text(
            "placeholder,description\n,空占位符\n{{ a }},字段A\n{{ a }},重复\n",
            encoding="utf-8",
        )

        with pytest.raises(TemplateGenError):
            import_placeholder_csv(str(source_csv), str(out_json))

    def test_import_cli_creates_output_parent_directory(self, tmp_path: Path, monkeypatch):
        source_csv = tmp_path / "edited.csv"
        out_json = tmp_path / "nested" / "dir" / "placeholder_descriptions.json"
        source_csv.write_text("placeholder,description\n{{ a }},字段A\n", encoding="utf-8")

        monkeypatch.setattr(
            "sys.argv",
            [
                "template_gen.import_placeholder_csv",
                "--input",
                str(source_csv),
                "--output",
                str(out_json),
            ],
        )
        import_main()

        assert out_json.exists()
