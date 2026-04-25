import pytest
import os
import tempfile
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docxtpl import DocxTemplate

from template_gen.parser import parse_document, print_table_coordinates
from template_gen.exceptions import ParseError
from template_gen.schemas import DocumentStructure, ParagraphInfo, TableInfo


class TestToolChainValidation:
    def test_parse_document_extracts_paragraphs(self):
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            doc = Document()
            doc.add_paragraph("First paragraph")
            doc.add_paragraph("Second paragraph")
            doc.save(f.name)
            temp_path = f.name
        
        structure = parse_document(temp_path)
        
        assert isinstance(structure, DocumentStructure)
        assert len(structure.paragraphs) >= 2
        assert any(p.text == "First paragraph" for p in structure.paragraphs)
        assert any(p.text == "Second paragraph" for p in structure.paragraphs)
        
        os.unlink(temp_path)
    
    def test_parse_document_extracts_tables(self):
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            doc = Document()
            table = doc.add_table(rows=2, cols=2)
            table.rows[0].cells[0].text = "Header A"
            table.rows[0].cells[1].text = "Header B"
            table.rows[1].cells[0].text = ""
            table.rows[1].cells[1].text = ""
            doc.save(f.name)
            temp_path = f.name
        
        structure = parse_document(temp_path)
        
        assert len(structure.tables) == 1
        assert len(structure.tables[0].rows) == 2
        assert structure.tables[0].rows[0].cells[0].text == "Header A"
        assert structure.tables[0].rows[1].cells[0].is_empty
        
        os.unlink(temp_path)
    
    def test_parse_document_raises_parse_error_for_invalid_file(self):
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            f.write(b"invalid content")
            temp_path = f.name
        
        with pytest.raises(ParseError):
            parse_document(temp_path)
        
        os.unlink(temp_path)
    
    def test_parse_document_raises_parse_error_for_missing_file(self):
        with pytest.raises(ParseError):
            parse_document("/nonexistent/path.docx")
    
    def test_docxtpl_fills_template_correctly(self):
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            doc = Document()
            para = doc.add_paragraph()
            run = para.add_run("{{ project_name }}")
            doc.save(f.name)
            template_path = f.name
        
        output_path = template_path.replace(".docx", "_filled.docx")
        
        tpl = DocxTemplate(template_path)
        tpl.render({"project_name": "Test Project"})
        tpl.save(output_path)
        
        filled_doc = Document(output_path)
        assert any("Test Project" in p.text for p in filled_doc.paragraphs)
        
        os.unlink(template_path)
        os.unlink(output_path)


class TestPrintTableCoordinates:
    def test_print_table_coordinates_outputs_structure(self):
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            doc = Document()
            table = doc.add_table(rows=2, cols=2)
            table.rows[0].cells[0].text = "Label"
            table.rows[0].cells[1].text = ""
            table.rows[1].cells[0].text = ""
            table.rows[1].cells[1].text = "Value"
            doc.save(f.name)
            temp_path = f.name
        
        output = print_table_coordinates(temp_path)
        
        assert "Table 0:" in output
        assert "[0,0]" in output
        assert "[0,1]" in output
        assert "empty" in output
        
        os.unlink(temp_path)