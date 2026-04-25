import pytest
import os
import tempfile
from docx import Document
from docx.shared import Pt

from template_gen.filler import detect_merged_cells, generate_template
from template_gen.generate_template import load_placeholders_json, parse_location
from template_gen.parser import parse_document
from template_gen.exceptions import TemplateGenError
from template_gen.schemas import CoordinateMapping


class TestGenerateTemplate:
    def test_parse_location_supports_paragraph_location(self):
        assert parse_location("paragraphs[0]") == {"paragraphs": 0}

    def test_parse_location_rejects_unsupported_locations(self):
        with pytest.raises(TemplateGenError):
            parse_location("tables[0].rows[0]")

    def test_load_placeholders_json_rejects_invalid_location(self):
        with tempfile.NamedTemporaryFile(suffix=".json", delete=False) as f:
            f.write(b'{"placeholders": [{"location": "tables[0].rows[bad].cells[1]", "placeholder": "{{ name }}"}]}')
            json_path = f.name

        with pytest.raises(TemplateGenError):
            load_placeholders_json(json_path)

        os.unlink(json_path)

    def test_load_placeholders_json_supports_paragraph_location(self):
        with tempfile.NamedTemporaryFile(suffix=".json", delete=False) as f:
            f.write(b'{"placeholders": [{"location": "paragraphs[1]", "placeholder": "{{ title }}"}]}')
            json_path = f.name

        mappings = load_placeholders_json(json_path)

        assert len(mappings) == 1
        assert mappings[0].target_type == "paragraph"
        assert mappings[0].paragraph_index == 1
        assert mappings[0].placeholder == "{{ title }}"

        os.unlink(json_path)

    def test_generate_template_inserts_placeholders_in_empty_cells(self):
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            doc = Document()
            table = doc.add_table(rows=3, cols=2)
            table.rows[0].cells[0].text = "Field A:"
            table.rows[0].cells[1].text = ""
            table.rows[1].cells[0].text = "Field B:"
            table.rows[1].cells[1].text = ""
            table.rows[2].cells[0].text = "Field C:"
            table.rows[2].cells[1].text = ""
            doc.save(f.name)
            input_path = f.name
        
        output_path = input_path.replace(".docx", "_template.docx")
        
        mappings = [
            CoordinateMapping(table_index=0, row_index=0, col_index=1, placeholder="{{ field_a }}", is_empty=True),
            CoordinateMapping(table_index=0, row_index=1, col_index=1, placeholder="{{ field_b }}", is_empty=True),
            CoordinateMapping(table_index=0, row_index=2, col_index=1, placeholder="{{ field_c }}", is_empty=True),
        ]
        
        generate_template(input_path, output_path, mappings)
        
        template_doc = Document(output_path)
        table = template_doc.tables[0]
        
        assert table.rows[0].cells[1].paragraphs[0].text == "{{ field_a }}"
        assert table.rows[1].cells[1].paragraphs[0].text == "{{ field_b }}"
        assert table.rows[2].cells[1].paragraphs[0].text == "{{ field_c }}"
        
        os.unlink(input_path)
        os.unlink(output_path)
    
    def test_generate_template_skips_non_empty_cells(self):
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            doc = Document()
            table = doc.add_table(rows=2, cols=2)
            table.rows[0].cells[0].text = "Label:"
            table.rows[0].cells[1].text = "Pre-filled Value"
            table.rows[1].cells[0].text = "Name:"
            table.rows[1].cells[1].text = ""
            doc.save(f.name)
            input_path = f.name
        
        output_path = input_path.replace(".docx", "_template.docx")
        
        mappings = [
            CoordinateMapping(table_index=0, row_index=0, col_index=1, placeholder="{{ should_skip }}", is_empty=False),
            CoordinateMapping(table_index=0, row_index=1, col_index=1, placeholder="{{ name }}", is_empty=True),
        ]
        
        generate_template(input_path, output_path, mappings)
        
        template_doc = Document(output_path)
        table = template_doc.tables[0]
        
        assert table.rows[0].cells[1].paragraphs[0].text == "Pre-filled Value"
        assert table.rows[1].cells[1].paragraphs[0].text == "{{ name }}"
        
        os.unlink(input_path)
        os.unlink(output_path)

    def test_generate_template_rechecks_actual_cell_emptiness(self):
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            doc = Document()
            table = doc.add_table(rows=1, cols=1)
            table.rows[0].cells[0].text = "Already filled"
            doc.save(f.name)
            input_path = f.name

        output_path = input_path.replace(".docx", "_template.docx")

        generate_template(
            input_path,
            output_path,
            [CoordinateMapping(table_index=0, row_index=0, col_index=0, placeholder="{{ should_not_write }}", is_empty=True)],
        )

        assert Document(output_path).tables[0].rows[0].cells[0].paragraphs[0].text == "Already filled"

        os.unlink(input_path)
        os.unlink(output_path)
    
    def test_generate_template_preserves_cell_styles(self):
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            doc = Document()
            table = doc.add_table(rows=1, cols=1)
            cell = table.rows[0].cells[0]
            para = cell.paragraphs[0]
            run = para.add_run("")
            run.font.name = "Times New Roman"
            run.font.size = Pt(14)
            run.bold = True
            doc.save(f.name)
            input_path = f.name
        
        output_path = input_path.replace(".docx", "_template.docx")
        
        mappings = [
            CoordinateMapping(table_index=0, row_index=0, col_index=0, placeholder="{{ styled_field }}", is_empty=True),
        ]
        
        generate_template(input_path, output_path, mappings)
        
        template_doc = Document(output_path)
        cell = template_doc.tables[0].rows[0].cells[0]
        run = cell.paragraphs[0].runs[0]
        
        assert run.text == "{{ styled_field }}"
        assert run.font.name == "Times New Roman"
        assert run.font.size == Pt(14)
        assert run.bold

        os.unlink(input_path)
        os.unlink(output_path)

    def test_generate_template_borrows_style_for_empty_cells_without_runs(self):
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            doc = Document()
            table = doc.add_table(rows=1, cols=2)
            label_run = table.rows[0].cells[0].paragraphs[0].add_run("项目名称")
            label_run.font.name = "仿宋"
            label_run.font.size = Pt(14)
            label_run.bold = True
            doc.save(f.name)
            input_path = f.name

        output_path = input_path.replace(".docx", "_template.docx")

        generate_template(
            input_path,
            output_path,
            [CoordinateMapping(table_index=0, row_index=0, col_index=1, placeholder="{{ project_name }}", is_empty=True)],
        )

        run = Document(output_path).tables[0].rows[0].cells[1].paragraphs[0].runs[0]

        assert run.text == "{{ project_name }}"
        assert run.font.name == "仿宋"
        assert run.font.size == Pt(14)
        assert run.bold

        os.unlink(input_path)
        os.unlink(output_path)

    def test_generate_template_supports_paragraph_placeholders(self):
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            doc = Document()
            para = doc.add_paragraph()
            run = para.add_run("")
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)
            run.bold = True
            doc.save(f.name)
            input_path = f.name

        output_path = input_path.replace(".docx", "_template.docx")

        generate_template(
            input_path,
            output_path,
            [CoordinateMapping(placeholder="{{ paragraph_field }}", is_empty=True, target_type="paragraph", paragraph_index=0)],
        )

        out_run = Document(output_path).paragraphs[0].runs[0]
        assert out_run.text == "{{ paragraph_field }}"
        assert out_run.font.name == "Times New Roman"
        assert out_run.font.size == Pt(12)
        assert out_run.bold

        os.unlink(input_path)
        os.unlink(output_path)

    def test_detect_merged_cells_handles_vertical_merges(self):
        doc = Document()
        table = doc.add_table(rows=3, cols=2)
        table.cell(0, 0).merge(table.cell(1, 0))

        merged_map = detect_merged_cells(table)

        assert merged_map[(1, 0)] == (0, 0)
