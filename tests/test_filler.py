import pytest
import os
import tempfile
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from template_gen.filler import batch_fill, fill_template, generate_template, set_cell_text_keep_basic_style
from template_gen.schemas import CoordinateMapping
from template_gen.exceptions import FillError


class TestSetCellTextKeepBasicStyle:
    def test_preserves_existing_font_style(self):
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            doc = Document()
            table = doc.add_table(rows=1, cols=1)
            cell = table.rows[0].cells[0]
            para = cell.paragraphs[0]
            run = para.add_run("")
            run.font.name = "Arial"
            run.font.size = Pt(12)
            doc.save(f.name)
            temp_path = f.name
        
        doc = Document(temp_path)
        cell = doc.tables[0].rows[0].cells[0]
        
        set_cell_text_keep_basic_style(cell, "{{ test_placeholder }}")
        
        doc.save(temp_path)
        saved_doc = Document(temp_path)
        saved_cell = saved_doc.tables[0].rows[0].cells[0]
        
        assert saved_cell.paragraphs[0].runs[0].text == "{{ test_placeholder }}"
        assert saved_cell.paragraphs[0].runs[0].font.name == "Arial"
        
        os.unlink(temp_path)
    
    def test_creates_run_when_none_exists(self):
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            doc = Document()
            table = doc.add_table(rows=1, cols=1)
            cell = table.rows[0].cells[0]
            cell.paragraphs[0].clear()
            doc.save(f.name)
            temp_path = f.name
        
        doc = Document(temp_path)
        cell = doc.tables[0].rows[0].cells[0]
        
        set_cell_text_keep_basic_style(cell, "{{ new_placeholder }}")
        
        doc.save(temp_path)
        saved_doc = Document(temp_path)
        saved_cell = saved_doc.tables[0].rows[0].cells[0]
        
        assert len(saved_cell.paragraphs[0].runs) == 1
        assert saved_cell.paragraphs[0].runs[0].text == "{{ new_placeholder }}"
        
        os.unlink(temp_path)
    
    def test_clears_secondary_runs(self):
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            doc = Document()
            table = doc.add_table(rows=1, cols=1)
            cell = table.rows[0].cells[0]
            para = cell.paragraphs[0]
            para.add_run("first")
            para.add_run("second")
            para.add_run("third")
            doc.save(f.name)
            temp_path = f.name
        
        doc = Document(temp_path)
        cell = doc.tables[0].rows[0].cells[0]
        
        set_cell_text_keep_basic_style(cell, "{{ replacement }}")
        
        doc.save(temp_path)
        saved_doc = Document(temp_path)
        saved_cell = saved_doc.tables[0].rows[0].cells[0]
        
        assert saved_cell.paragraphs[0].runs[0].text == "{{ replacement }}"
        assert len(saved_cell.paragraphs[0].runs) == 3
        assert saved_cell.paragraphs[0].runs[1].text == ""
        assert saved_cell.paragraphs[0].runs[2].text == ""
        
        os.unlink(temp_path)


class TestFillTemplate:
    def test_fill_template_raises_error_for_missing_required_data(self):
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            doc = Document()
            doc.add_paragraph("{{ project_name }}")
            doc.save(f.name)
            template_path = f.name

        output_path = template_path.replace(".docx", "_filled.docx")

        with pytest.raises(FillError):
            fill_template(template_path, {}, output_path)

        os.unlink(template_path)

    def test_fill_template_preserves_style_for_generated_empty_cell_placeholders(self):
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            doc = Document()
            table = doc.add_table(rows=1, cols=2)
            label_run = table.rows[0].cells[0].paragraphs[0].add_run("项目名称")
            label_run.font.name = "仿宋"
            label_run.font.size = Pt(14)
            label_run.bold = True
            doc.save(f.name)
            input_path = f.name

        template_path = input_path.replace(".docx", "_template.docx")
        output_path = input_path.replace(".docx", "_filled.docx")

        generate_template(
            input_path,
            template_path,
            [CoordinateMapping(table_index=0, row_index=0, col_index=1, placeholder="{{ project_name }}", is_empty=True)],
        )

        fill_template(template_path, {"project_name": "人工智能培训"}, output_path)

        filled_run = Document(output_path).tables[0].rows[0].cells[1].paragraphs[0].runs[0]
        assert filled_run.text == "人工智能培训"
        assert filled_run.font.name == "仿宋"
        assert filled_run.font.size == Pt(14)
        assert filled_run.bold

        os.unlink(input_path)
        os.unlink(template_path)
        os.unlink(output_path)

    def test_fill_template_replaces_placeholders(self):
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            doc = Document()
            para = doc.add_paragraph()
            para.add_run("{{ project_name }}")
            table = doc.add_table(rows=1, cols=1)
            table.rows[0].cells[0].paragraphs[0].add_run("{{ applicant_name }}")
            doc.save(f.name)
            template_path = f.name
        
        output_path = template_path.replace(".docx", "_filled.docx")
        
        result = fill_template(template_path, {
            "project_name": "My Project",
            "applicant_name": "John Doe"
        }, output_path)
        
        assert result == output_path
        
        filled_doc = Document(output_path)
        texts = [p.text for p in filled_doc.paragraphs]
        assert "My Project" in texts
        assert filled_doc.tables[0].rows[0].cells[0].paragraphs[0].text == "John Doe"
        
        os.unlink(template_path)
        os.unlink(output_path)
    
    def test_fill_template_raises_error_for_missing_template(self):
        with pytest.raises(FillError):
            fill_template("/nonexistent/template.docx", {}, "/tmp/output.docx")


class TestBatchFill:
    def test_batch_fill_creates_multiple_documents(self):
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            doc = Document()
            para = doc.add_paragraph()
            para.add_run("{{ name }}")
            doc.save(f.name)
            template_path = f.name
        
        output_dir = tempfile.mkdtemp()
        
        records = [
            {"name": "Document1"},
            {"name": "Document2"},
            {"name": "Document3"},
        ]
        
        outputs = batch_fill(template_path, records, output_dir)
        
        assert len(outputs) == 3
        
        for i, path in enumerate(outputs):
            filled_doc = Document(path)
            assert records[i]["name"] in filled_doc.paragraphs[0].text
            os.unlink(path)
        
        os.unlink(template_path)
        os.rmdir(output_dir)
    
    def test_batch_fill_with_custom_filename_pattern(self):
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            doc = Document()
            para = doc.add_paragraph()
            para.add_run("{{ name }}")
            doc.save(f.name)
            template_path = f.name
        
        output_dir = tempfile.mkdtemp()
        
        records = [
            {"name": "Alpha"},
            {"name": "Beta"},
        ]
        
        outputs = batch_fill(
            template_path,
            records,
            output_dir,
            filename_pattern="{name}.docx"
        )
        
        assert outputs[0].endswith("Alpha.docx")
        assert outputs[1].endswith("Beta.docx")
        
        for path in outputs:
            os.unlink(path)
        
        os.unlink(template_path)
        os.rmdir(output_dir)

    def test_batch_fill_sanitizes_filename_placeholders(self):
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            doc = Document()
            para = doc.add_paragraph()
            para.add_run("{{ name }}")
            doc.save(f.name)
            template_path = f.name

        output_dir = tempfile.mkdtemp()

        outputs = batch_fill(
            template_path,
            [{"name": "../../escape"}],
            output_dir,
            filename_pattern="{name}.docx"
        )

        assert len(outputs) == 1
        assert os.path.abspath(outputs[0]).startswith(os.path.abspath(output_dir) + os.sep)
        assert outputs[0].endswith("__escape.docx")

        os.unlink(outputs[0])
        os.unlink(template_path)
        os.rmdir(output_dir)
    
    def test_batch_fill_handles_partial_data(self):
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            doc = Document()
            table = doc.add_table(rows=3, cols=2)
            table.rows[0].cells[0].text = "Course 1:"
            table.rows[0].cells[1].paragraphs[0].add_run("{{ course_1_name }}")
            table.rows[1].cells[0].text = "Course 2:"
            table.rows[1].cells[1].paragraphs[0].add_run("{{ course_2_name }}")
            table.rows[2].cells[0].text = "Course 3:"
            table.rows[2].cells[1].paragraphs[0].add_run("{{ course_3_name }}")
            doc.save(f.name)
            template_path = f.name
        
        output_dir = tempfile.mkdtemp()
        
        records = [
            {"course_1_name": "Math", "course_2_name": "Science", "course_3_name": ""},
        ]
        
        outputs = batch_fill(template_path, records, output_dir)
        
        filled_doc = Document(outputs[0])
        cells = filled_doc.tables[0].rows
        
        assert cells[0].cells[1].paragraphs[0].text == "Math"
        assert cells[1].cells[1].paragraphs[0].text == "Science"
        assert cells[2].cells[1].paragraphs[0].text == ""
        
        os.unlink(outputs[0])
        os.unlink(template_path)
        os.rmdir(output_dir)
